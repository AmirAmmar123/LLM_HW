import re
import os
from docx import Document
from multiprocessing import Pool, cpu_count
import json



PATH = r"./knesset_protocols"


INVALID_TALKERS_NAMES = {
    'חברי כנסת',
    'משתתפים באמצעים מקוונים',
    'מנהלות הוועדה',
    'מסקנות הוועדה'
    "חברי הוועדה",
    "מוזמנים",
    "קריאה",
    "אני מבקש לשאול",
    'נכחו',
    'סדר היום',
    'חברי הכנסת',
    'באמצעים מקוונים',
    'ייעוץ משפטי',
    'באמצעים מקוונים',
    'רישום פרלמנטרי',
    'באמצעים מקוונים',
    '',
    'חברי הוועדה',
    '<הישיבה ננעלה בשעה 11',
    'כתוב במלווה',
    'לגבי ההצעה של להוסיף',
    'ב- יש טענה אחרת כתוב',
    'אז לכתוב',
    '<ייעוץ משפטי',
    'הישיבה ננעלה בשעה 12',
    '<הישיבה ננעלה בשעה 09',
    'אני רק רוצה לתזכר',
    'מנהל/ת הוועדה',
    'אני אגיד עוד פעם',
    'בואו נעשה סדר רגע',
    '- - אמרו',
    'אני רוצה לשאול ספציפית',
    'אז יש לי שאלה'
}

GET_RID_OF_SUFFIX ={
    'יו"ר ועדת הכנסת',
    'שר החינוך',
    'שר החינוך '
    'היו"ר',
    'היו”ר',
    'לאיכות הסביבה'
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה'
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'סגנית מזכיר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',
    'היו"ר',
    'היו”ר',
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה',
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'השר',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',

    # Ministers
    'שר הביטחון',
    'שר החוץ',
    'שר האנרגיה',
    'שר האנרגיה והתשתיות',
    'שר המדע',
    'שר הרווחה',
    'שר הכלכלה',
    'שר הקליטה',
    'שר הדתות',
    'שר לביטחון הפנים',
    'שר לביטחון פנים',
    'שר המים',

    'השרה להגנת הסביבה',
    'השרה לחינוך',
    'השרה למדע',
    'השרה לחדשנות',
    'השרה לשוויון חברתי',
    'שרת הקליטה',
    'שרת התחבורה',
    'שרת הבריאות',
    'שרת החינוך',
    'שרת התיירות',
    'שרת התרבות והספורט',

    # Deputy ministers
    'סגן שר',
    'סגנית שר',
    'סגן שר הביטחון',
    'סגן שר האוצר',
    'סגן שר החוץ',
    'סגן שר הבריאות',
    'סגן שר התחבורה',
    'סגן שר המשפטים',
    'סגן שר החינוך',
    'ממלא מקום שר',
    'מ״מ שר',

    # Knesset roles
    'יושב ראש הכנסת',
    'יו״ר הכנסת',
    'יו"ר הכנסת',
    'סגן יושב ראש הכנסת',
    'סגן יו״ר הכנסת',
    'סגנית יו״ר הכנסת',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'סגנית מזכיר הכנסת',
    'מנהל הוועדה',
    'מנהלת הוועדה',
    'מזכיר הוועדה',
    'מזכירת הוועדה',
    'יועץ משפטי',
    'יועצת משפטית',
    'היועץ המשפטי',
    'היועצת המשפטית',
    'יועמ״ש הכנסת',
    'יועמ"ש הוועדה',

    # MK roles & political positions
    'ח"כ',
    'ח״כ',
    'יו״ר הסיעה',
    'יושב ראש הסיעה',
    'ראש האופוזיציה',
    'ראש הקואליציה',
    'רכז הקואליציה',
    'רכז האופוזיציה',
    'מזכ״ל המפלגה',
    'מ״מ יו״ר',

    # Government & authority roles
    'מנכ״ל המשרד',
    'מנכ"ל המשרד',
    'מנכ״ל',
    'מנכ"ל',
    'סמנכ״ל',
    'סמנכ"ל',
    'המשנה למנכ״ל',
    'המשנה למנכ"ל',
    'ראש הרשות',
    'מפקד מחוז',
    'קצין אגף',
    'ראש אגף',
    'סגן ראש אגף',
    'ממונה על',
    'ממונה תחום',
    'מנהל תחום',
    'מנהל מחלקה',
    'ראש מחלקה',

    # External roles that appear often
    'מבקר המדינה',
    'נציב שירות המדינה',
    'נציב הכבאות',
    'דובר הכנסת',
    'מבקר הכנסת',
    'מבקר המשרד',

    # General labels to remove
    'מוזמנים',
    'מוזמן',
    'אורחים',
    'נוכחו',
    'נכחו',
    'רשם הישיבה',
    'השר',
    'השרה',
    'החינוך, התרבות והספורט',
    'התרבות והספורט',
    ',',
    'מזכירת הכנסת',
    'לאיכות הסביבה',
    'לאזרחים ותיקים',
    'שר המודיעין',
    'ופיתוח הכפר',
    'ראש הממשלה',
    'שר התשתיות',
    'שר העבודה והרווחה',
    'סגן',
    'סגנית',
    'שרת המשפטים',
    'ופיתוח הכפר',
    'עו"ד',
    'הלאומיות',
    'שר התשתיות',
    '(יו"ר הוועדה המסדרת)',

}


HEBREW_UNITS = {
    "אפס": 0,
    "אחד": 1, "אחת": 1,
    "שניים": 2, "שתיים": 2, "שני": 2, "שתי": 2,
    "שלוש": 3, "שלושה": 3,
    "ארבע": 4, "ארבעה": 4,
    "חמש": 5, "חמישה": 5,
    "שש": 6, "שישה": 6,
    "שבע": 7, "שבעה": 7,
    "שמונה": 8,
    "תשע": 9, "תשעה": 9
}

HEBREW_TENS = {
    "עשר": 10, "עשרה": 10,
    "עשרים": 20,
    "שלושים": 30,
    "ארבעים": 40,
    "חמישים": 50,
    "שישים": 60,
    "שבעים": 70,
    "שמונים": 80,
    "תשעים": 90,
}

HEBROW_HUNDREDS = {
    "מאה": 100,
    "מאתיים": 200,
    "שלוש מאות": 300,
    "שלוש־מאות": 300,   
    "ארבע מאות": 400,
    "חמש מאות": 500,
    "שש מאות": 600,
    "שבע מאות": 700,
    "שמונה מאות": 800,
    "תשע מאות": 900
}



START_WITH = ['סדר', 'הישיבה','חברי', 'ייעוץ', 'מנהלי', 'רישום', 'הודעה','.לא','אני','אמרתי', 'אדוני', 'הביטחון','אבל','מדובר', 'הנושא','דקה','מובילה','לקריאה', 'יש','אוקיי',
              'א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י', 'כ', 'ל', 'מ', 'נ', 'ס', 'ע', 'פ', 'צ', 'ק', 'ר', 'ש', 'ת', 'ועוד', 'התשובה', 'ואני', 'נעבור', 'לכן', 'אולי', 
              'בכלל','והטכנולוגיה','הצעת','קריאת','שאלתי','במלים','אנחנו', 'זו', 'עכשיו', 'להעביר', 'הצעת', 'ובכלל', 'לכן', 'נעבור', 'בואו', 'תיכף', 'בורג', 'התבררו', 'אז', 'בפרשת', 
              'הבנתי', 'דבר', 'על', 'וברשותך', 'או', 'אסיים', 'בסדר.', 'זה','לא', 'אגיד', 'אתן', 'אלו', 'הנקודה','הצעה', 'ולכן', 'קוראת', 'הוא', 'האם', 'העניין', 'הדברים', 'הדבר', 'היום', 'הטענה',
              'הכוונה', 'הכרעה', 'החלטה', 'החלטה', 'היא', 'היא', 'הייתי', 'הייתם', 'הייתן', 'היית', 'לא.','לגבי', 'הסעיף', 'אם','החוזר', 'שאלתי', 'מנהל/ת', 'מנהל', 'השאלת', 'דוגמה', 'והמשפט'
              'ס–התש"ס–2000;', 'בוודאי.', 'וצריך', 'הנה']


def process_file(filename):
    """Process a single file and return a dictionary with protocol info or None."""
    result = FileParser(filename).parse_filename()
    if result:
        knesset_number, protocol_type = result
        protocol = Protocol(knesset_number, protocol_type, filename)
        if protocol.protocol_number is not None:
            return {
                "knesset_number": protocol.knesset_number,
                "protocol_type": protocol.protocol_type,
                "protocol_number": protocol.protocol_number,
                "chair": protocol.yor_hankest,
                "filename": filename,
                "Speakers": protocol.speakers,
                "map_sentence": protocol.map_sentence
            }
    else:
        print(f"Filename: {filename} => Invalid format")
    return None



class FileLoader:
    """Class to load files from a specified directory."""
    def __init__(self, path:str = PATH):
        self.path = path
    
    def ListFiles(self):
        return os.listdir(self.path)


class FileParser:
    """Class to parse filenames and extract metadata."""

    def __init__(self, filename:str):
        self.filename = filename    

        
    def parse_filename(self):
        if not self.filename:
            return None
        match = re.match(r'^(\d{1,2})_pt([mv])_.*\.docx$', self.filename)
        if not match:
            return None

        knesset_number = int(match.group(1))
        letter = match.group(2)
        protocol_type = "plenary" if letter == 'm' else "committee"

        return knesset_number, protocol_type


class Protocol:
    def __init__(self, knesset_number: int, protocol_type: str, filename: str):
        self.knesset_number = knesset_number
        self.protocol_type = protocol_type
        self.filepath = os.path.join(PATH, filename)
        self.protocol_number = self._extract_protocol_number()
        if self.protocol_number is None:
            return
        
        self.yor_hankest = self._extract_yor()
        self.speakers = self._extract_speakers_names()
        self.map_sentence = self._extract_speakers_sentences()

    def _extract_speakers_sentences(self):
        """
        Map each extracted speaker to the sentences they said in the document.
        Uses the cleaned list of speakers from _extract_speakers_names().
        """
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for sentences: {e}")
            return {}

        speaker_sentences = {}
        current_speaker = None

        # Get the list of cleaned speakers
        speakers_list = self._extract_speakers_names()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph contains a colon
            if ':' in text:
                potential_speaker = text.split(":", 1)[0].strip()
                # Clean potential speaker name similarly to _extract_speakers_names
                for suffix in GET_RID_OF_SUFFIX:
                    potential_speaker = potential_speaker.replace(suffix, "").strip()
                potential_speaker = re.sub(r'\([^)]*\)', '', potential_speaker).strip()
                potential_speaker = " ".join(potential_speaker.split())

                # Check if the cleaned name matches one of the extracted speakers
                if potential_speaker in speakers_list:
                    current_speaker = potential_speaker
                    if current_speaker not in speaker_sentences:
                        speaker_sentences[current_speaker] = []

                # Add the text after the colon as the first sentence
                if current_speaker:
                    sentence = text.split(":", 1)[1].strip()
                    if sentence:
                        speaker_sentences[current_speaker].append(sentence)
            else:
                # Continuation of previous speaker
                if current_speaker:
                    speaker_sentences[current_speaker].append(text)

        return speaker_sentences



    def _extract_protocol_number(self):
        """Extract the protocol number from the document."""
        try:
            self.doc = Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath}: {e}")
            return None

        text = "\n".join(p.text for p in self.doc.paragraphs)

 
        m1 = re.search(r"פרוטוקול\s+מס'?[\s:]*([0-9]+)", text)
    
        m2 = re.search(r"הישיבה\s+([א-ת\s\-־]+)", text)

        if m1:
            return int(m1.group(1))
        if m2:
            return self._hebrew_number_to_int(m2.group(1))

        return -1
    

    def _hebrew_number_to_int(self, text):
        """Convert Hebrew number words into integers"""

        text = text.replace("־", " ").replace("-", " ").replace("–", " ")
        words = text.split()


        if "של" in words:
            words = words[:words.index("של")]

   
        if text.strip() == "אלף":
            return 1000

        total = 0
        i = 0
        while i < len(words):
            w = words[i]

        
            while w.startswith(("ה", "ו")) and len(w) > 1:
                w = w[1:]

  
            if w == "ו":
                i += 1
                continue

      
            if i + 1 < len(words):
                next_w = words[i+1].lstrip("ו").lstrip("ה")
                pair = f"{w} {next_w}"
                if pair in HEBROW_HUNDREDS:
                    total += HEBROW_HUNDREDS[pair]
                    i += 2
                    continue


            if w in HEBROW_HUNDREDS:
                total += HEBROW_HUNDREDS[w]
            elif w in HEBREW_TENS:
                total += HEBREW_TENS[w]
            elif w in HEBREW_UNITS:
                total += HEBREW_UNITS[w]

            i += 1

        return total if total > 0 else -1

    
    def _extract_yor(self):
        """
        Extract the chairperson's name using multiple patterns:

        Priority:
        1) 'חברי הוועדה: <name> – היו"ר'
        2) היו"ר <name>:
        3) << >> / < > cases
        4) יו"ר / יו״ר / יור formats
        5) fallback: first speaker with colon
        """

     
        text = "\n".join(p.text.strip() for p in self.doc.paragraphs if p.text.strip())

  
        text_clean = re.sub(r'[<>{}]+', ' ', text)
        text_clean = re.sub(r'\s+', ' ', text_clean)


        opening_pattern = (
            r'דברי\s+פתיחה[:\s]+(?:חבר(?:ת)?\s+הכנסת\s+)?'
            r'([א-ת\"\'״׳\(\)\s\-]+?)\s*[-–]\s*יו"?ר'
        )

        m = re.search(opening_pattern, text_clean)
        if m:
            return m.group(1).strip()
        
        committee_pattern = (
            r'חברי הוועדה[:\s]+([א-ת\"\'״׳\(\)\s\-]+?)\s*[–—\-]\s*(?:היו"ר|יו"ר|יו״ר|יור)'
        )

        m = re.search(committee_pattern, text_clean)
        if m:
            return m.group(1).strip()

        chair_patterns = [
            r'היו"ר\s+([א-ת\"\'״׳\.\-\s]+?):',
            r'יו"ר\s+([א-ת\"\'״׳\.\-\s]+?):',
            r'יו״ר\s+([א-ת\"\'״׳\.\-\s]+?):',
            r'יור\s+([א-ת\"\'״׳\.\-\s]+?):',
        ]

        for pat in chair_patterns:
            m = re.search(pat, text_clean)
            if m:
                return m.group(1).strip()

        for line in text.split("\n"):
            if ":" in line:
                clean = re.sub(r'[<>{}]+', '', line).strip()
                if any(x in clean for x in ['היו"ר', 'יו"ר', 'יו״ר', 'יור']):
                    left = clean.split(":")[0]
                    for x in ['היו"ר', 'יו"ר', 'יו״ר', 'יור']:
                        left = left.replace(x, "")
                    return left.strip()

        return "Unknown"

    def extract_text_between_markers(self,text):
        """
        Extract text from a line:
        - If <something>, return 'something'
        - If <<s1>> something <<s2>>, return 'something'
        """
        text = text.strip()
        

        double_match = re.match(r'^<{2}.*?>{2}\s*(.*?)\s*<{2}.*?}>{2}$', text)
        if double_match:
            return double_match.group(1).strip()

        single_match = re.match(r'^<{1,2}\s*(.*?)\s*>{1,2}$', text)
        if single_match:
            return single_match.group(1).strip()
        
        return None
    

    def _has_markers(self, text):
        return re.findall(r'<{1,2}.*?>{1,2}', text)


    def _parse_double_markers(self, text, markers):
        first_end = text.find(markers[0]) + len(markers[0])
        last_start = text.rfind(markers[-1])
        return text[first_end:last_start].strip()
    

    def _extract_speakers_names(self):
        """
        Extract paragraphs that contain ':' and clean potential speaker names.
        Handles <something> and <<s1>> something <<s2>> correctly.
        """
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for colon sentences: {e}")
            return []

        results = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if ':' not in text:
                continue

     
            markers = self._has_markers(text)

            if markers:
        
                if len(markers) == 2:
                    clean_text = self._parse_double_markers(text, markers)
                else:
                    clean_text = self.extract_text_between_markers(markers[0])

                if not clean_text:
                    continue

                if clean_text in INVALID_TALKERS_NAMES:
                    continue

                for suffix in GET_RID_OF_SUFFIX:
                    clean_text = clean_text.replace(suffix, "").strip()

         
                clean_text = re.sub(r'\([^)]*\)', '', clean_text).strip()

                clean_text = clean_text.rstrip(':').strip()

                clean_text = " ".join(clean_text.split())

            
                parts = clean_text.split()
                if parts and  parts[0] in START_WITH:
                    continue

                if 2 <= len(parts) <= 3:
                    results.append(clean_text)
                continue  


            clean_text = text.split(":", 1)[0].strip()
            if clean_text in INVALID_TALKERS_NAMES:
                continue

            for suffix in GET_RID_OF_SUFFIX:
                clean_text = clean_text.replace(suffix, "").strip()

            clean_text = re.sub(r'\([^)]*\)', '', clean_text).strip()
    
            parts = clean_text.split()

            if parts and parts[0] in START_WITH:
                continue

            if 2 <= len(parts) <= 3:
                results.append(clean_text)


            
       
        unique = []
        seen = set()
        for s in results:
            if 5 < len(s) < 50 and s not in seen:
                seen.add(s)
                unique.append(s)


        if not unique:
            with open("debug_files.txt", "a", encoding="utf-8") as f:
                f.write(self.filepath + "\n")

        return unique

        

class ProtocolsCollection:
    def __init__(self):
        self.protocols = []

    def add_protocol(self, protocol):
        self.protocols.append(protocol)


# debug_files = [
#         "16_ptv_577443.docx",
#         # "16_ptv_577758.docx",
#         # "16_ptv_491962.docx",
#         # "20_ptv_490139.docx",
#         # "15_ptv_498215.docx",
# ]


if __name__ == "__main__":
    fl = FileLoader()
    files = fl.ListFiles()


    with Pool(processes=cpu_count()) as pool:
        protocols = pool.map(process_file, files)

    # protocols = [process_file(f) for f in debug_files ]
    # protocols = [p for p in protocols if p is not None]


    for p in protocols:
        if p:
            print(f"Knesset: {p['knesset_number']}, Type: {p['protocol_type']}, Protocol Number: {p['protocol_number']}, Chair: {p['chair']}, file: {p['filename']}")
            print("Speakers and sentences:")
            for speaker, sentences in p['map_sentence'].items():
                print(f"\nSpeaker: {speaker}")
                for idx, sentence in enumerate(sentences, 1):
                    print(f"  {idx}. {sentence}")
            print("\n" + "="*80 + "\n")