import re
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from multiprocessing import Pool, cpu_count
import json



PATH = r"./knesset_protocols"


INVALID_TALKERS_NAMES = {
    'חברי כנסת',
    'משתתפים באמצעים מקוונים',
    'מנהלות הוועדה',
    'מסקנות הוועדה'
    "חברי הוועדה",
    "מוזמנים",
    "קריאה",
    "אני מבקש לשאול",
    'נכחו',
    'סדר היום',
    'חברי הכנסת',
    'באמצעים מקוונים',
    'ייעוץ משפטי',
    'באמצעים מקוונים',
    'רישום פרלמנטרי',
    'באמצעים מקוונים',
    '',
    'חברי הוועדה'

}

GET_RID_OF_SUFFIX ={
    'יו"ר ועדת הכנסת',
    'שר החינוך',
    'שר החינוך '
    'היו"ר',
    'היו”ר',
    'לאיכות הסביבה'
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה'
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'סגנית מזכיר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',
    'היו"ר',
    'היו”ר',
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה',
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'השר',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',

    # Ministers
    'שר הביטחון',
    'שר החוץ',
    'שר האנרגיה',
    'שר האנרגיה והתשתיות',
    'שר המדע',
    'שר הרווחה',
    'שר הכלכלה',
    'שר הקליטה',
    'שר הדתות',
    'שר לביטחון הפנים',
    'שר לביטחון פנים',
    'שר המים',

    'השרה להגנת הסביבה',
    'השרה לחינוך',
    'השרה למדע',
    'השרה לחדשנות',
    'השרה לשוויון חברתי',
    'שרת הקליטה',
    'שרת התחבורה',
    'שרת הבריאות',
    'שרת החינוך',
    'שרת התיירות',
    'שרת התרבות והספורט',

    # Deputy ministers
    'סגן שר',
    'סגנית שר',
    'סגן שר הביטחון',
    'סגן שר האוצר',
    'סגן שר החוץ',
    'סגן שר הבריאות',
    'סגן שר התחבורה',
    'סגן שר המשפטים',
    'סגן שר החינוך',
    'ממלא מקום שר',
    'מ״מ שר',

    # Knesset roles
    'יושב ראש הכנסת',
    'יו״ר הכנסת',
    'יו"ר הכנסת',
    'סגן יושב ראש הכנסת',
    'סגן יו״ר הכנסת',
    'סגנית יו״ר הכנסת',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'סגנית מזכיר הכנסת',
    'מנהל הוועדה',
    'מנהלת הוועדה',
    'מזכיר הוועדה',
    'מזכירת הוועדה',
    'יועץ משפטי',
    'יועצת משפטית',
    'היועץ המשפטי',
    'היועצת המשפטית',
    'יועמ״ש הכנסת',
    'יועמ"ש הוועדה',

    # MK roles & political positions
    'ח"כ',
    'ח״כ',
    'יו״ר הסיעה',
    'יושב ראש הסיעה',
    'ראש האופוזיציה',
    'ראש הקואליציה',
    'רכז הקואליציה',
    'רכז האופוזיציה',
    'מזכ״ל המפלגה',
    'מ״מ יו״ר',

    # Government & authority roles
    'מנכ״ל המשרד',
    'מנכ"ל המשרד',
    'מנכ״ל',
    'מנכ"ל',
    'סמנכ״ל',
    'סמנכ"ל',
    'המשנה למנכ״ל',
    'המשנה למנכ"ל',
    'ראש הרשות',
    'מפקד מחוז',
    'קצין אגף',
    'ראש אגף',
    'סגן ראש אגף',
    'ממונה על',
    'ממונה תחום',
    'מנהל תחום',
    'מנהל מחלקה',
    'ראש מחלקה',

    # External roles that appear often
    'מבקר המדינה',
    'נציב שירות המדינה',
    'נציב הכבאות',
    'דובר הכנסת',
    'מבקר הכנסת',
    'מבקר המשרד',

    # General labels to remove
    'מוזמנים',
    'מוזמן',
    'אורחים',
    'נוכחו',
    'נכחו',
    'רשם הישיבה',
    'השר',
    'השרה',
    'החינוך, התרבות והספורט',
    'התרבות והספורט',
    ',',
    'מזכירת הכנסת',
    'לאיכות הסביבה',
    'לאזרחים ותיקים',
    'שר המודיעין',
    'ופיתוח הכפר',
    'ראש הממשלה',
    'שר התשתיות',
    'שר העבודה והרווחה',
    'סגן',
    'סגנית',
    'שרת המשפטים',
    'ופיתוח הכפר',
    'עו"ד',
    'הלאומיות',
    'שר התשתיות'

}


HEBREW_UNITS = {
    "אפס": 0,
    "אחד": 1, "אחת": 1,
    "שניים": 2, "שתיים": 2, "שני": 2, "שתי": 2,
    "שלוש": 3, "שלושה": 3,
    "ארבע": 4, "ארבעה": 4,
    "חמש": 5, "חמישה": 5,
    "שש": 6, "שישה": 6,
    "שבע": 7, "שבעה": 7,
    "שמונה": 8,
    "תשע": 9, "תשעה": 9
}

HEBREW_TENS = {
    "עשר": 10, "עשרה": 10,
    "עשרים": 20,
    "שלושים": 30,
    "ארבעים": 40,
    "חמישים": 50,
    "שישים": 60,
    "שבעים": 70,
    "שמונים": 80,
    "תשעים": 90,
}

HEBROW_HUNDREDS = {
    "מאה": 100,
    "מאתיים": 200,
    "שלוש מאות": 300,
    "שלוש־מאות": 300,   
    "ארבע מאות": 400,
    "חמש מאות": 500,
    "שש מאות": 600,
    "שבע מאות": 700,
    "שמונה מאות": 800,
    "תשע מאות": 900
}


def process_file(filename):
    """Process a single file and return a dictionary with protocol info or None."""
    result = FileParser(filename).parse_filename()
    if result:
        knesset_number, protocol_type = result
        protocol = Protocol(knesset_number, protocol_type, filename)
        if protocol.protocol_number is not None:
            return {
                "knesset_number": protocol.knesset_number,
                "protocol_type": protocol.protocol_type,
                "protocol_number": protocol.protocol_number,
                "chair": protocol.yor_hankest,
                "filename": filename,
                "Speakers": protocol.colon_sentences,
                "Speeches": protocol.speeches
            }
    else:
        print(f"Filename: {filename} => Invalid format")
    return None



class FileLoader:
    """Class to load files from a specified directory."""
    def __init__(self, path:str = PATH):
        self.path = path
    
    def ListFiles(self):
        return os.listdir(self.path)


class FileParser:
    """Class to parse filenames and extract metadata."""

    def __init__(self, filename:str):
        self.filename = filename    

        
    def parse_filename(self):
        if not self.filename:
            return None
        match = re.match(r'^(\d{1,2})_pt([mv])_.*\.docx$', self.filename)
        if not match:
            return None

        knesset_number = int(match.group(1))
        letter = match.group(2)
        protocol_type = "plenary" if letter == 'm' else "committee"

        return knesset_number, protocol_type


class Protocol:
    def __init__(self, knesset_number: int, protocol_type: str, filename: str):
        self.knesset_number = knesset_number
        self.protocol_type = protocol_type
        self.filepath = os.path.join(PATH, filename)
        self.protocol_number = self._extract_protocol_number()
        if self.protocol_number is None:
            return
        
        self.yor_hankest = self._extract_yor()
        self.colon_sentences = self._extract_colon_sentences()
        self.speeches = self._extract_speeches()

    def _extract_protocol_number(self):
        """Extract the protocol number from the document."""
        try:
            self.doc = Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath}: {e}")
            return None

        text = "\n".join(p.text for p in self.doc.paragraphs)

 
        m1 = re.search(r"פרוטוקול\s+מס'?[\s:]*([0-9]+)", text)
    
        m2 = re.search(r"הישיבה\s+([א-ת\s\-־]+)", text)

        if m1:
            return int(m1.group(1))
        if m2:
            return self._hebrew_number_to_int(m2.group(1))

        return -1
    

    def _hebrew_number_to_int(self, text):
        """Convert Hebrew number words into integers"""

        text = text.replace("־", " ").replace("-", " ").replace("–", " ")
        words = text.split()


        if "של" in words:
            words = words[:words.index("של")]

   
        if text.strip() == "אלף":
            return 1000

        total = 0
        i = 0
        while i < len(words):
            w = words[i]

        
            while w.startswith(("ה", "ו")) and len(w) > 1:
                w = w[1:]

  
            if w == "ו":
                i += 1
                continue

      
            if i + 1 < len(words):
                next_w = words[i+1].lstrip("ו").lstrip("ה")
                pair = f"{w} {next_w}"
                if pair in HEBROW_HUNDREDS:
                    total += HEBROW_HUNDREDS[pair]
                    i += 2
                    continue


            if w in HEBROW_HUNDREDS:
                total += HEBROW_HUNDREDS[w]
            elif w in HEBREW_TENS:
                total += HEBREW_TENS[w]
            elif w in HEBREW_UNITS:
                total += HEBREW_UNITS[w]

            i += 1

        return total if total > 0 else -1

    
    def _extract_yor(self):
        """Extract only the first chairperson's full name from the document."""
        text = "\n".join(p.text for p in self.doc.paragraphs)


        match = re.search(r'היו"ר\s+([א-ת\"\'״׳\.\-\s]+?)(:|\n|מוזמנים|$)', text)
        if match:
            full_name = match.group(1).strip()

            for line in full_name.splitlines():
                line = line.strip()
                if line:
                    return line
        return "Unknown"


    def _normalize_speaker_from_text(self, text: str):
        """Return cleaned speaker name if text represents a speaker line."""
        candidate = text.strip()
        if not candidate.endswith(':'):
            return None

        candidate = candidate.rstrip(':').strip()
        if candidate in INVALID_TALKERS_NAMES:
            return None

        clean_text = candidate
        for suffix in GET_RID_OF_SUFFIX:
            clean_text = clean_text.replace(suffix, "").strip()

        clean_text = re.sub(r'\([^)]*\)', '', clean_text).strip()
        clean_text = " ".join(clean_text.split())
        if not clean_text or clean_text in INVALID_TALKERS_NAMES:
            return None

        parts = [p.strip() for p in clean_text.split(' ') if p.strip()]
        if len(parts) < 2 or len(parts) > 5:
            return None

        return " ".join(parts)

    def _extract_colon_sentences(self):
        """
        Extract paragraphs that:
        - end with a colon ':'
        - contain at least one run that is bold OR underlined
        """
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for colon sentences: {e}")
            return []

        results = []
        for para in doc.paragraphs:
            if self._is_heading_paragraph(para):
                continue
            text = para.text.strip()
            speaker_name = self._normalize_speaker_from_text(text)
            if speaker_name:
                results.append(speaker_name)

        seen = set()
        unique = []
        for s in results:
            if 5 < len(s) < 50 and s not in seen:
                seen.add(s)
                unique.append(s)

        return unique

    def _extract_speeches(self):
        """Attach textual content to the speakers identified in the document."""
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for speeches: {e}")
            return []

        speeches = []
        current_speaker = None
        current_chunks = []

        def flush_current():
            if current_speaker and current_chunks:
                full_text = " ".join(current_chunks).strip()
                speeches.append({
                    "speaker": current_speaker,
                    "text": full_text,
                    "sentences": self._split_into_sentences(full_text)
                })

        for para in doc.paragraphs:
            if self._is_heading_paragraph(para):
                continue
            text = para.text.strip()
            if not text:
                continue

            speaker_name = self._normalize_speaker_from_text(text)
            if speaker_name:
                flush_current()
                current_speaker = speaker_name
                current_chunks = []
                continue

            if current_speaker:
                current_chunks.append(text)

        flush_current()
        return speeches

    def _is_single_token_enum(self, text: str, idx: int):
        """Detect enumerations like 'א.' or '1.' to avoid sentence splits."""
        j = idx - 1
        while j >= 0 and text[j].isspace():
            j -= 1
        if j < 0:
            return False
        start = j
        while start - 1 >= 0 and text[start - 1].isalpha():
            start -= 1
        token = text[start:j + 1]
        if len(token) == 1 and (start == 0 or text[start - 1].isspace() or text[start - 1] in '(["\''):
            return True
        if token.isdigit():
            if start == 0 or text[start - 1].isspace() or text[start - 1] in '(["\'':
                return True
        return False

    def _split_into_sentences(self, text: str):
        """Split a text block into crude sentences (no filtering)."""
        if not text:
            return []

        normalized = text.replace('\r', ' ').replace('\n', ' ')
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        if not normalized:
            return []

        sentences = []
        buffer = []
        length = len(normalized)

        sentence_endings = {'.', '!'}

        i = 0
        while i < length:
            ch = normalized[i]
            buffer.append(ch)

            if ch in sentence_endings:
                prev_char = normalized[i - 1] if i > 0 else ''
                next_char = normalized[i + 1] if i + 1 < length else ''

                if prev_char.isdigit() and next_char.isdigit():
                    i += 1
                    continue

                if ch == '.' and self._is_single_token_enum(normalized, i):
                    i += 1
                    continue

                sentence = ''.join(buffer).strip()
                if sentence:
                    sentences.append(self._clean_sentence(sentence))
                buffer = []

            i += 1

        tail = ''.join(buffer).strip()
        if tail:
            sentences.append(self._clean_sentence(tail))

        return sentences

    def _is_heading_paragraph(self, para):
        """Detect paragraphs that are likely headings (centered/bold labels)."""
        text = para.text.strip()
        if not text:
            return True
        if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        runs = [run for run in para.runs if run.text.strip()]
        if runs and all(run.bold for run in runs):
            return True
        return False

    def _clean_sentence(self, sentence: str):
        """Remove English letters and collapse repeated symbols."""
        if not sentence:
            return sentence

        cleaned = re.sub(r'[A-Za-z]', '', sentence)
        cleaned = re.sub(r'(?:-\s*){2,}', ' ', cleaned)
        cleaned = re.sub(r'-{2,}', ' ', cleaned)
        cleaned = re.sub(r'-\s*$', '', cleaned)
        cleaned = re.sub(r'([^\w\s\u0590-\u05FF-])\1+', r'\1', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned


        

class ProtocolsCollection:
    def __init__(self):
        self.protocols = []

    def add_protocol(self, protocol):
        self.protocols.append(protocol)


import re
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from multiprocessing import Pool, cpu_count
import json



PATH = r"./knesset_protocols"


INVALID_TALKERS_NAMES = {
    'חברי כנסת',
    'משתתפים באמצעים מקוונים',
    'מנהלות הוועדה',
    'מסקנות הוועדה'
    "חברי הוועדה",
    "מוזמנים",
    "קריאה",
    "אני מבקש לשאול",
    'נכחו',
    'סדר היום',
    'חברי הכנסת',
    'באמצעים מקוונים',
    'ייעוץ משפטי',
    'באמצעים מקוונים',
    'רישום פרלמנטרי',
    'באמצעים מקוונים',
    '',
    'חברי הוועדה'

}

GET_RID_OF_SUFFIX ={
    'יו"ר ועדת הכנסת',
    'שר החינוך',
    'שר החינוך '
    'היו"ר',
    'היו”ר',
    'לאיכות הסביבה'
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה'
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'סגנית מזכיר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',
    'היו"ר',
    'היו”ר',
    'השר לאיכות הסביבה',
    'שר הבריאות',
    'סגן שר החינוך, התרבות והספורט',
    'שר החינוך, התרבות והספורט',
    'שר המשפטים',
    'רצוני לשאול',
    'שר האוצר',
    'שרת העבודה והרווחה',
    'שר התחבורה',
    'תשובת שר התחבורה',
    'שר המשטרה',
    'תשובת שר המשטרה',
    'סגנית מזכיר הכנסת',
    'שר החקלאות',
    'חבר הכנסת',
    'שאל את ביום',
    'משרד המשפטים',
    'שר התיירות',
    'סגן מזכיר הכנסת',
    'השר',
    'שר התקשורת',
    'השר לענייני דתות',
    'מזכיר הכנסת',
    'מנהלת הוועדה',
    'שר הפנים',
    'תשובת',
    'להגנת הסביבה',
    'השר להגנת הסביבה',
    'שר התשתיות הלאומיות',

    # Ministers
    'שר הביטחון',
    'שר החוץ',
    'שר האנרגיה',
    'שר האנרגיה והתשתיות',
    'שר המדע',
    'שר הרווחה',
    'שר הכלכלה',
    'שר הקליטה',
    'שר הדתות',
    'שר לביטחון הפנים',
    'שר לביטחון פנים',
    'שר המים',

    'השרה להגנת הסביבה',
    'השרה לחינוך',
    'השרה למדע',
    'השרה לחדשנות',
    'השרה לשוויון חברתי',
    'שרת הקליטה',
    'שרת התחבורה',
    'שרת הבריאות',
    'שרת החינוך',
    'שרת התיירות',
    'שרת התרבות והספורט',

    # Deputy ministers
    'סגן שר',
    'סגנית שר',
    'סגן שר הביטחון',
    'סגן שר האוצר',
    'סגן שר החוץ',
    'סגן שר הבריאות',
    'סגן שר התחבורה',
    'סגן שר המשפטים',
    'סגן שר החינוך',
    'ממלא מקום שר',
    'מ״מ שר',

    # Knesset roles
    'יושב ראש הכנסת',
    'יו״ר הכנסת',
    'יו"ר הכנסת',
    'סגן יושב ראש הכנסת',
    'סגן יו״ר הכנסת',
    'סגנית יו״ר הכנסת',
    'מזכיר הכנסת',
    'סגן מזכיר הכנסת',
    'סגנית מזכיר הכנסת',
    'מנהל הוועדה',
    'מנהלת הוועדה',
    'מזכיר הוועדה',
    'מזכירת הוועדה',
    'יועץ משפטי',
    'יועצת משפטית',
    'היועץ המשפטי',
    'היועצת המשפטית',
    'יועמ״ש הכנסת',
    'יועמ"ש הוועדה',

    # MK roles & political positions
    'ח"כ',
    'ח״כ',
    'יו״ר הסיעה',
    'יושב ראש הסיעה',
    'ראש האופוזיציה',
    'ראש הקואליציה',
    'רכז הקואליציה',
    'רכז האופוזיציה',
    'מזכ״ל המפלגה',
    'מ״מ יו״ר',

    # Government & authority roles
    'מנכ״ל המשרד',
    'מנכ"ל המשרד',
    'מנכ״ל',
    'מנכ"ל',
    'סמנכ״ל',
    'סמנכ"ל',
    'המשנה למנכ״ל',
    'המשנה למנכ"ל',
    'ראש הרשות',
    'מפקד מחוז',
    'קצין אגף',
    'ראש אגף',
    'סגן ראש אגף',
    'ממונה על',
    'ממונה תחום',
    'מנהל תחום',
    'מנהל מחלקה',
    'ראש מחלקה',

    # External roles that appear often
    'מבקר המדינה',
    'נציב שירות המדינה',
    'נציב הכבאות',
    'דובר הכנסת',
    'מבקר הכנסת',
    'מבקר המשרד',

    # General labels to remove
    'מוזמנים',
    'מוזמן',
    'אורחים',
    'נוכחו',
    'נכחו',
    'רשם הישיבה',
    'השר',
    'השרה',
    'החינוך, התרבות והספורט',
    'התרבות והספורט',
    ',',
    'מזכירת הכנסת',
    'לאיכות הסביבה',
    'לאזרחים ותיקים',
    'שר המודיעין',
    'ופיתוח הכפר',
    'ראש הממשלה',
    'שר התשתיות',
    'שר העבודה והרווחה',
    'סגן',
    'סגנית',
    'שרת המשפטים',
    'ופיתוח הכפר',
    'עו"ד',
    'הלאומיות',
    'שר התשתיות'

}


HEBREW_UNITS = {
    "אפס": 0,
    "אחד": 1, "אחת": 1,
    "שניים": 2, "שתיים": 2, "שני": 2, "שתי": 2,
    "שלוש": 3, "שלושה": 3,
    "ארבע": 4, "ארבעה": 4,
    "חמש": 5, "חמישה": 5,
    "שש": 6, "שישה": 6,
    "שבע": 7, "שבעה": 7,
    "שמונה": 8,
    "תשע": 9, "תשעה": 9
}

HEBREW_TENS = {
    "עשר": 10, "עשרה": 10,
    "עשרים": 20,
    "שלושים": 30,
    "ארבעים": 40,
    "חמישים": 50,
    "שישים": 60,
    "שבעים": 70,
    "שמונים": 80,
    "תשעים": 90,
}

HEBROW_HUNDREDS = {
    "מאה": 100,
    "מאתיים": 200,
    "שלוש מאות": 300,
    "שלוש־מאות": 300,   
    "ארבע מאות": 400,
    "חמש מאות": 500,
    "שש מאות": 600,
    "שבע מאות": 700,
    "שמונה מאות": 800,
    "תשע מאות": 900
}


def process_file(filename):
    """Process a single file and return a dictionary with protocol info or None."""
    result = FileParser(filename).parse_filename()
    if result:
        knesset_number, protocol_type = result
        protocol = Protocol(knesset_number, protocol_type, filename)
        if protocol.protocol_number is not None:
            return {
                "knesset_number": protocol.knesset_number,
                "protocol_type": protocol.protocol_type,
                "protocol_number": protocol.protocol_number,
                "chair": protocol.yor_hankest,
                "filename": filename,
                "Speakers": protocol.colon_sentences,
                "Speeches": protocol.speeches
            }
    else:
        print(f"Filename: {filename} => Invalid format")
    return None



class FileLoader:
    """Class to load files from a specified directory."""
    def __init__(self, path:str = PATH):
        self.path = path
    
    def ListFiles(self):
        return os.listdir(self.path)


class FileParser:
    """Class to parse filenames and extract metadata."""

    def __init__(self, filename:str):
        self.filename = filename    

        
    def parse_filename(self):
        if not self.filename:
            return None
        match = re.match(r'^(\d{1,2})_pt([mv])_.*\.docx$', self.filename)
        if not match:
            return None

        knesset_number = int(match.group(1))
        letter = match.group(2)
        protocol_type = "plenary" if letter == 'm' else "committee"

        return knesset_number, protocol_type


class Protocol:
    def __init__(self, knesset_number: int, protocol_type: str, filename: str):
        self.knesset_number = knesset_number
        self.protocol_type = protocol_type
        self.filepath = os.path.join(PATH, filename)
        self.protocol_number = self._extract_protocol_number()
        if self.protocol_number is None:
            return
        
        self.yor_hankest = self._extract_yor()
        self.colon_sentences = self._extract_colon_sentences()
        self.speeches = self._extract_speeches()

    def _extract_protocol_number(self):
        """Extract the protocol number from the document."""
        try:
            self.doc = Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath}: {e}")
            return None

        text = "\n".join(p.text for p in self.doc.paragraphs)

 
        m1 = re.search(r"פרוטוקול\s+מס'?[\s:]*([0-9]+)", text)
    
        m2 = re.search(r"הישיבה\s+([א-ת\s\-־]+)", text)

        if m1:
            return int(m1.group(1))
        if m2:
            return self._hebrew_number_to_int(m2.group(1))

        return -1
    

    def _hebrew_number_to_int(self, text):
        """Convert Hebrew number words into integers"""

        text = text.replace("־", " ").replace("-", " ").replace("–", " ")
        words = text.split()


        if "של" in words:
            words = words[:words.index("של")]

   
        if text.strip() == "אלף":
            return 1000

        total = 0
        i = 0
        while i < len(words):
            w = words[i]

        
            while w.startswith(("ה", "ו")) and len(w) > 1:
                w = w[1:]

  
            if w == "ו":
                i += 1
                continue

      
            if i + 1 < len(words):
                next_w = words[i+1].lstrip("ו").lstrip("ה")
                pair = f"{w} {next_w}"
                if pair in HEBROW_HUNDREDS:
                    total += HEBROW_HUNDREDS[pair]
                    i += 2
                    continue


            if w in HEBROW_HUNDREDS:
                total += HEBROW_HUNDREDS[w]
            elif w in HEBREW_TENS:
                total += HEBREW_TENS[w]
            elif w in HEBREW_UNITS:
                total += HEBREW_UNITS[w]

            i += 1

        return total if total > 0 else -1

    
    def _extract_yor(self):
        """Extract only the first chairperson's full name from the document."""
        text = "\n".join(p.text for p in self.doc.paragraphs)


        match = re.search(r'היו"ר\s+([א-ת\"\'״׳\.\-\s]+?)(:|\n|מוזמנים|$)', text)
        if match:
            full_name = match.group(1).strip()

            for line in full_name.splitlines():
                line = line.strip()
                if line:
                    return line
        return "Unknown"


    def _normalize_speaker_from_text(self, text: str):
        """Return cleaned speaker name if text represents a speaker line."""
        candidate = text.strip()
        if not candidate.endswith(':'):
            return None

        candidate = candidate.rstrip(':').strip()
        if candidate in INVALID_TALKERS_NAMES:
            return None

        clean_text = candidate
        for suffix in GET_RID_OF_SUFFIX:
            clean_text = clean_text.replace(suffix, "").strip()

        clean_text = re.sub(r'\([^)]*\)', '', clean_text).strip()
        clean_text = " ".join(clean_text.split())
        if not clean_text or clean_text in INVALID_TALKERS_NAMES:
            return None

        parts = [p.strip() for p in clean_text.split(' ') if p.strip()]
        if len(parts) < 2 or len(parts) > 5:
            return None

        return " ".join(parts)

    def _extract_colon_sentences(self):
        """
        Extract paragraphs that:
        - end with a colon ':'
        - contain at least one run that is bold OR underlined
        """
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for colon sentences: {e}")
            return []

        results = []
        for para in doc.paragraphs:
            if self._is_heading_paragraph(para):
                continue
            text = para.text.strip()
            speaker_name = self._normalize_speaker_from_text(text)
            if speaker_name:
                results.append(speaker_name)

        seen = set()
        unique = []
        for s in results:
            if 5 < len(s) < 50 and s not in seen:
                seen.add(s)
                unique.append(s)

        return unique

    def _extract_speeches(self):
        """Attach textual content to the speakers identified in the document."""
        try:
            doc = self.doc if hasattr(self, "doc") else Document(self.filepath)
        except Exception as e:
            print(f"Error opening {self.filepath} for speeches: {e}")
            return []

        speeches = []
        current_speaker = None
        current_chunks = []

        def flush_current():
            if current_speaker and current_chunks:
                full_text = " ".join(current_chunks).strip()
                speeches.append({
                    "speaker": current_speaker,
                    "text": full_text,
                    "sentences": self._split_into_sentences(full_text)
                })

        for para in doc.paragraphs:
            if self._is_heading_paragraph(para):
                continue
            text = para.text.strip()
            if not text:
                continue

            speaker_name = self._normalize_speaker_from_text(text)
            if speaker_name:
                flush_current()
                current_speaker = speaker_name
                current_chunks = []
                continue

            if current_speaker:
                current_chunks.append(text)

        flush_current()
        return speeches

    def _is_single_token_enum(self, text: str, idx: int):
        """Detect enumerations like 'א.' or '1.' to avoid sentence splits."""
        j = idx - 1
        while j >= 0 and text[j].isspace():
            j -= 1
        if j < 0:
            return False
        start = j
        while start - 1 >= 0 and text[start - 1].isalpha():
            start -= 1
        token = text[start:j + 1]
        if len(token) == 1 and (start == 0 or text[start - 1].isspace() or text[start - 1] in '(["\''):
            return True
        if token.isdigit():
            if start == 0 or text[start - 1].isspace() or text[start - 1] in '(["\'':
                return True
        return False

    def _split_into_sentences(self, text: str):
        """Split a text block into crude sentences (no filtering)."""
        if not text:
            return []

        normalized = text.replace('\r', ' ').replace('\n', ' ')
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        if not normalized:
            return []

        sentences = []
        buffer = []
        length = len(normalized)

        sentence_endings = {'.', '!'}

        i = 0
        while i < length:
            ch = normalized[i]
            buffer.append(ch)

            if ch in sentence_endings:
                prev_char = normalized[i - 1] if i > 0 else ''
                next_char = normalized[i + 1] if i + 1 < length else ''

                if prev_char.isdigit() and next_char.isdigit():
                    i += 1
                    continue

                if ch == '.' and self._is_single_token_enum(normalized, i):
                    i += 1
                    continue

                sentence = ''.join(buffer).strip()
                if sentence:
                    sentences.append(self._clean_sentence(sentence))
                buffer = []

            i += 1

        tail = ''.join(buffer).strip()
        if tail:
            sentences.append(self._clean_sentence(tail))

        return sentences

    def _is_heading_paragraph(self, para):
        """Detect paragraphs that are likely headings (centered/bold labels)."""
        text = para.text.strip()
        if not text:
            return True
        if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        runs = [run for run in para.runs if run.text.strip()]
        if runs and all(run.bold for run in runs):
            return True
        return False

    def _clean_sentence(self, sentence: str):
        """Remove English letters and collapse repeated symbols."""
        if not sentence:
            return sentence

        cleaned = re.sub(r'[A-Za-z]', '', sentence)
        cleaned = re.sub(r'(?:-\s*){2,}', ' ', cleaned)
        cleaned = re.sub(r'-{2,}', ' ', cleaned)
        cleaned = re.sub(r'-\s*$', '', cleaned)
        cleaned = re.sub(r'([^\w\s\u0590-\u05FF-])\1+', r'\1', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned


        

class ProtocolsCollection:
    def __init__(self):
        self.protocols = []

    def add_protocol(self, protocol):
        self.protocols.append(protocol)

if __name__ == "__main__":
    fl = FileLoader()
    files = fl.ListFiles()
    with Pool(processes=cpu_count()) as pool:
        results = pool.map(process_file, files)
    protocols = [p for p in results if p]
    output_file = "protocols_output.jsonl"
    with open(output_file, "w", encoding="utf-8") as f:
        for p in protocols:
            protocol_name = p.get("filename", "")
            knesset_number = p.get("knesset_number", "")
            protocol_type = p.get("protocol_type", "")
            protocol_number = p.get("protocol_number", "")
            protocol_chairman = p.get("chair", "")
            speeches = p.get("Speeches", [])  # Use [] as default for list
            for speech in speeches:
                speaker = speech["speaker"]
                sentences = speech["sentences"]
                for sentence in sentences:
                    line = {
                        "protocol_name": protocol_name,
                        "knesset_number": knesset_number,
                        "protocol_type": protocol_type,
                        "protocol_number": protocol_number,
                        "protocol_chairman": protocol_chairman,
                        "speaker_name": speaker,
                        "sentence_text": sentence
                    }
                    f.write(json.dumps(line, ensure_ascii=False) + "\n")
    print(f"\nJSONL file saved to {output_file}")